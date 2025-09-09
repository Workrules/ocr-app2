import streamlit as st
import pypdfium2 as pdfium
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
import os
import io
from PIL import Image, ImageFile
import cv2
import numpy as np
import json
from openai import OpenAI
import re
import difflib
from itertools import zip_longest
from typing import List, Tuple, Dict, Any
import gc
import time

# === Word出力関連 ===
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==== 環境変数 ====
AZURE_ENDPOINT = os.getenv("AZURE_DOCINT_ENDPOINT")
AZURE_KEY = os.getenv("AZURE_DOCINT_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
MODEL_NAME = os.getenv("OCR_GPT_MODEL", "gpt-5")  # gpt-5 / gpt-5-mini など
BATCH_SIZE = max(1, int(os.getenv("OCR_BATCH_PAGES", "10")))  # 10ページずつ処理

# ==== 共有辞書の場所（環境変数で指定） ====
DICT_DIR = os.getenv("OCR_DICT_DIR", ".")
DICT_FILE = os.path.join(DICT_DIR, "ocr_char_corrections.json")
UNTRAINED_FILE = os.path.join(DICT_DIR, "untrained_confusions.json")
TRAINED_FILE = os.path.join(DICT_DIR, "trained_confusions.json")

# ==== 事前チェック ====
if not AZURE_ENDPOINT or not AZURE_KEY:
    st.error("環境変数 AZURE_DOCINT_ENDPOINT と AZURE_DOCINT_KEY を設定してください。")
    st.stop()

# ==== クライアント初期化 ====
client = DocumentAnalysisClient(endpoint=AZURE_ENDPOINT, credential=AzureKeyCredential(AZURE_KEY))
openai_client = OpenAI(api_key=OPENAI_API_KEY)

# ==== 正規表現 ====
JP_CHAR_RE = re.compile(r"^[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF]$")

# ==== ユーティリティ ====
def remove_red_stamp(img_pil: Image.Image) -> Image.Image:
    img = np.array(img_pil)
    hsv = cv2.cvtColor(img, cv2.COLOR_RGB2HSV)
    lower_red1 = np.array([0, 70, 50]); upper_red1 = np.array([10, 255, 255])
    lower_red2 = np.array([170, 70, 50]); upper_red2 = np.array([180, 255, 255])
    mask = cv2.bitwise_or(cv2.inRange(hsv, lower_red1, upper_red1),
                          cv2.inRange(hsv, lower_red2, upper_red2))
    img[mask > 0] = [255, 255, 255]
    return Image.fromarray(img)

def load_json(path: str, retries: int = 3, delay: float = 0.1) -> dict:
    for _ in range(retries):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except FileNotFoundError:
            return {}
        except (json.JSONDecodeError, PermissionError):
            time.sleep(delay)
    return {}

def save_json(obj: dict, path: str):
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    tmp_path = f"{path}.tmp"
    with open(tmp_path, "w", encoding="utf-8") as f:
        json.dump(obj, f, ensure_ascii=False, indent=2)
        f.flush(); os.fsync(f.fileno())
    os.replace(tmp_path, path)

def learn_charwise_with_missing(original: str, corrected: str) -> dict:
    learned = {}
    sm = difflib.SequenceMatcher(None, original, corrected)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag in ["replace", "insert"]:
            o_seg = original[i1:i2]
            c_seg = corrected[j1:j2]
            for o_char, c_char in zip_longest(o_seg, c_seg, fillvalue=""):
                if c_char and (not o_char or o_char != c_char):
                    wrong = o_char if o_char else "□"
                    if JP_CHAR_RE.match(c_char) or c_char == "□":
                        learned[wrong] = {"right": c_char, "count": 1}
    return learned

def update_dictionary_and_untrained(learned: dict):
    dictionary = load_json(DICT_FILE)
    for w, meta in learned.items():
        if w in dictionary:
            if dictionary[w]["right"] == meta["right"]:
                dictionary[w]["count"] += meta["count"]
            else:
                if meta["count"] > dictionary[w]["count"]:
                    dictionary[w] = meta
        else:
            dictionary[w] = meta
    save_json(dictionary, DICT_FILE)

    untrained = load_json(UNTRAINED_FILE)
    for w, meta in learned.items():
        untrained[w] = meta["right"]
    save_json(untrained, UNTRAINED_FILE)

def gpt_fix_text(text: str, dictionary: dict) -> str:
    prompt = f"""
次のOCR結果を自然な日本語に直してください。
- 日本語に存在しない文字は「□」にしてください。
- 辞書候補を参考にしてください: {json.dumps(dictionary, ensure_ascii=False)}
- 意味を勝手に補完せず、最小限の修正だけ行ってください。

OCR結果:
{text}
""".strip()
    try:
        resp = openai_client.responses.create(
            model=MODEL_NAME,
            input=prompt,
            text={"verbosity": "low"},
            reasoning={"effort": "minimal"}
        )
        if hasattr(resp, "output_text") and resp.output_text:
            return resp.output_text.strip()
        out_parts = []
        for item in getattr(resp, "output", []) or []:
            for part in getattr(item, "content", []) or []:
                t = getattr(part, "text", None)
                if t:
                    out_parts.append(t)
        out = "".join(out_parts).strip()
        return out or text
    except Exception as e:
        st.warning(f"GPT補正をスキップしました（エラー）：{e}")
        return text

def is_pdf(b: bytes) -> bool:
    return len(b) >= 5 and b[:5] == b"%PDF-"

def render_pdf_selected_pages(pdf_bytes: bytes, indices_0based: List[int], dpi: int = 200) -> Tuple[List[Image.Image], List[int]]:
    imgs: List[Image.Image] = []; nums: List[int] = []
    pdf = pdfium.PdfDocument(io.BytesIO(pdf_bytes))
    scale = dpi / 72.0
    for idx in indices_0based:
        page = pdf[idx]
        pil = page.render(scale=scale).to_pil().convert("RGB")
        imgs.append(pil); nums.append(idx + 1)
    return imgs, nums

def parse_page_spec(spec: str, max_pages: int) -> List[int]:
    s = (spec or "").strip()
    if not s:
        return []
    out = set()
    parts = [p.strip() for p in s.split(",") if p.strip()]
    for p in parts:
        if "-" in p:
            a, b = p.split("-", 1)
            try:
                start = max(1, min(int(a), int(b)))
                end = min(max_pages, max(int(a), int(b)))
                for n in range(start, end + 1):
                    out.add(n - 1)
            except ValueError:
                continue
        else:
            try:
                n = int(p)
                if 1 <= n <= max_pages:
                    out.add(n - 1)
            except ValueError:
                continue
    return sorted(out)

def chunked(seq: List[int], n: int) -> List[List[int]]:
    return [seq[i:i+n] for i in range(0, len(seq), n)]

# === EMU / cm 変換ヘルパー ===
EMU_PER_CM = 360000.0
def to_cm(val) -> float:
    try:
        return float(getattr(val, "cm"))
    except Exception:
        try:
            return float(val) / EMU_PER_CM
        except Exception:
            return float(val)

# === 行のバウンディングボックス取得 ===
def line_bbox(line_obj):
    poly = getattr(line_obj, "polygon", None) or getattr(line_obj, "bounding_polygon", None)
    if not poly:
        return 0.0, 0.0, 0.0, 0.0
    xs, ys = [], []
    for p in poly:
        x = getattr(p, "x", None); y = getattr(p, "y", None)
        if x is None and isinstance(p, dict):
            x = p.get("x", 0.0); y = p.get("y", 0.0)
        xs.append(float(x)); ys.append(float(y))
    return min(xs or [0.0]), max(xs or [0.0]), min(ys or [0.0]), max(ys or [0.0])

# === インデント段のクラスタリング ===
def cluster_indents_cm(indent_values_cm, epsilon_cm=0.6):
    vals = sorted(indent_values_cm)
    clusters = []
    for v in vals:
        if not clusters:
            clusters.append({"center_cm": v, "members": [v]}); continue
        if abs(v - clusters[-1]["center_cm"]) <= epsilon_cm:
            c = clusters[-1]; c["members"].append(v)
            c["center_cm"] = sum(c["members"]) / len(c["members"])
        else:
            clusters.append({"center_cm": v, "members": [v]})
    return clusters

# === Word生成（整列/インデント段を反映） ===
def build_docx_from_layout(
    pages_layout: List[Dict[str, Any]],
    *,
    add_label: bool = False,
    center_thresh_ratio: float = 0.06,
    right_thresh_ratio: float = 0.04,
    indent_epsilon_cm: float = 0.6
) -> Tuple[bytes, List[Dict[str, Any]]]:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0); section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)

    page_w_cm = to_cm(section.page_width)
    left_cm = to_cm(section.left_margin); right_cm = to_cm(section.right_margin)
    content_width_cm = max(0.1, page_w_cm - left_cm - right_cm)

    style = doc.styles["Normal"]; style.font.name = "Yu Gothic"; style.font.size = Pt(11)

    all_indents_cm = []
    for page in pages_layout:
        pw = float(page.get("page_width") or 1.0)
        for ln in page.get("lines", []):
            indent_cm = (ln["x_min"] / max(pw, 1e-6)) * content_width_cm
            all_indents_cm.append(indent_cm)

    clusters = cluster_indents_cm(all_indents_cm, epsilon_cm=indent_epsilon_cm)
    clusters = sorted(clusters, key=lambda c: c["center_cm"])
    for idx, c in enumerate(clusters):
        c["index"] = idx; c["count"] = len(c["members"])

    def nearest_cluster_idx(indent_cm: float) -> int:
        if not clusters: return 0
        return min(range(len(clusters)), key=lambda i: abs(indent_cm - clusters[i]["center_cm"]))

    for page_i, page in enumerate(pages_layout, start=1):
        pw = float(page.get("page_width") or 1.0)
        ph = float(page.get("page_height") or 1.0)
        lines = sorted(page.get("lines", []), key=lambda r: (r["y"], r["x_min"]))

        center_px = pw / 2.0
        center_allow_px = pw * center_thresh_ratio
        right_allow_px = pw * right_thresh_ratio

        prev_y = None
        y_thresh = ph * 0.018

        for item in lines:
            txt = item["text"]; x_min = float(item["x_min"]); x_max = float(item["x_max"]); y = float(item["y"])
            para = doc.add_paragraph()

            line_center = (x_min + x_max) / 2.0
            align = "LEFT"
            if abs(line_center - center_px) <= center_allow_px:
                align = "CENTER"
            elif (max(0.0, pw - x_max)) <= right_allow_px:
                align = "RIGHT"

            indent_cm = (x_min / max(pw, 1e-6)) * content_width_cm
            ind_idx = nearest_cluster_idx(indent_cm)
            snapped_cm = clusters[ind_idx]["center_cm"] if clusters else indent_cm

            if add_label:
                tag = f"[{align} | ind={ind_idx}] "
                lab = para.add_run(tag); lab.font.size = Pt(8); lab.font.color.rgb = RGBColor(120, 120, 120)

            para.add_run(txt)
            para.paragraph_format.left_indent = Cm(max(0.0, min(0.95 * content_width_cm, snapped_cm)))

            if align == "CENTER":
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "RIGHT":
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            para.paragraph_format.space_after = Pt(2)
            if prev_y is not None and (y - prev_y) > y_thresh:
                para.paragraph_format.space_before = Pt(8)
            prev_y = y

        if page_i < len(pages_layout):
            doc.add_page_break()

    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    summary = [{"index": c["index"], "center_cm": round(c["center_cm"], 2), "count": c["count"]} for c in clusters]
    return bio.read(), summary

# ==== UI ====
st.title("📄 rouki-ocr（共有辞書） - ページ先指定 / 10ページバッチ / GPT補正 / Word出力（整列＆段）")

# 診断・設定
st.sidebar.subheader("📖 辞書の参照先")
st.sidebar.write({
    "OCR_DICT_DIR": os.path.abspath(DICT_DIR),
    "DICT_FILE": DICT_FILE,
    "UNTRAINED_FILE": UNTRAINED_FILE,
})
dictionary_preview = load_json(DICT_FILE)
st.sidebar.json(dictionary_preview)

st.sidebar.markdown("### 🔧 環境")
st.sidebar.write({
    "AZURE_DOCINT_ENDPOINT_set": bool(AZURE_ENDPOINT),
    "AZURE_DOCINT_KEY_set": bool(AZURE_KEY),
    "OPENAI_API_KEY_set": bool(OPENAI_API_KEY),
    "MODEL": MODEL_NAME,
    "BATCH_SIZE(default)": BATCH_SIZE,
})

# --- デバッグ項目 ---
st.sidebar.markdown("### 🛠 デバッグ")
skip_gpt = st.sidebar.checkbox("GPT補正をスキップ", value=False)
ocr_timeout = st.sidebar.slider("OCRタイムアウト（秒）", 10, 180, 60, step=5)
batch_size_override = st.sidebar.number_input("バッチサイズ上書き", 1, 20, value=BATCH_SIZE)

uploaded_file = st.file_uploader("画像またはPDFをアップロードしてください", type=["jpg", "jpeg", "png", "pdf"])
if not uploaded_file:
    st.info("📂 ここにファイルをアップロードしてください"); st.stop()

file_bytes = uploaded_file.read()
ImageFile.LOAD_TRUNCATED_IMAGES = True
is_input_pdf = uploaded_file.type == "application/pdf" or uploaded_file.name.lower().endswith(".pdf") or is_pdf(file_bytes)

# ==== PDF：OCR前にページ指定フォーム ====
if is_input_pdf:
    try:
        pdf_for_count = pdfium.PdfDocument(io.BytesIO(file_bytes))
        total_pages = len(pdf_for_count)
    except Exception as e:
        st.exception(e); st.stop()

    with st.form("pdf_select_form"):
        st.subheader("▶ OCRするページを先に選択")
        select_mode = st.radio("選択方法", options=["全ページ", "範囲指定", "ページ番号指定（例: 1,3,5-7）"],
                               index=1 if total_pages > 1 else 0, horizontal=True)
        dpi = st.slider("レンダリングDPI（高いほど精細・重い）", 72, 300, 200, step=4)

        center_pct = st.slider("中央揃え判定しきい値（%/ページ幅）", 2, 12, 6, step=1)
        indent_eps = st.slider("インデント段まとめ しきい値（cm）", 0.2, 1.5, 0.6, step=1/10)
        show_labels = st.checkbox("Wordにレイアウトラベルを書く（[CENTER | ind=n]）", value=False)

        if select_mode == "範囲指定" and total_pages > 1:
            start, end = st.slider("処理するページ範囲（1始まり）", 1, total_pages, (1, min(total_pages, 5)))
            chosen_indices = list(range(start - 1, end))
        elif select_mode == "ページ番号指定（例: 1,3,5-7）":
            spec = st.text_input("ページ番号（カンマ区切り、範囲はハイフン）",
                                 value="1-3" if total_pages >= 3 else "1")
            chosen_indices = parse_page_spec(spec, total_pages)
            if not chosen_indices:
                st.info("有効なページ番号を入力してください。例: 1,3,5-7")
        else:
            chosen_indices = list(range(total_pages))

        submitted = st.form_submit_button("このページだけOCRを実行")

    if not submitted or not chosen_indices:
        st.stop()

    EFFECTIVE_BATCH = int(batch_size_override) if batch_size_override else BATCH_SIZE

    total_to_process = len(chosen_indices)
    progress = st.progress(0.0); status = st.empty()
    all_corrected_texts: List[str] = []
    pages_layout: List[Dict[str, Any]] = []
    done = 0

    for batch_no, batch_indices in enumerate(chunked(chosen_indices, EFFECTIVE_BATCH), start=1):
        status.info(f"🔄 バッチ {batch_no} / {((total_to_process - 1) // EFFECTIVE_BATCH) + 1} を処理中（ページ: {', '.join(str(i+1) for i in batch_indices)}）")
        try:
            pages, page_numbers = render_pdf_selected_pages(file_bytes, batch_indices, dpi=dpi)
        except Exception as e:
            st.exception(e); st.stop()

        for page_img, page_num in zip(pages, page_numbers):
            st.write(f"## ページ {page_num}")
            clean_img = remove_red_stamp(page_img)

            buf = io.BytesIO(); clean_img.save(buf, format="PNG"); buf.seek(0)

            with st.spinner("OCRを実行中..."):
                t0 = time.perf_counter()
                try:
                    poller = client.begin_analyze_document("prebuilt-read", document=buf)
                    result = poller.result(timeout=float(ocr_timeout))
                except Exception as e:
                    st.error(f"OCRが{ocr_timeout}秒以内に完了しませんでした / 失敗しました：{e}")
                    st.caption(f"OCR実行時間: {time.perf_counter() - t0:.1f}s")
                    continue
                st.caption(f"OCR実行時間: {time.perf_counter() - t0:.1f}s")

            doc_page = result.pages[0] if getattr(result, "pages", None) else None
            if not doc_page:
                st.warning("OCR結果にページが見つかりませんでした。")
                done += 1; progress.progress(done / total_to_process); continue

            azure_lines = getattr(doc_page, "lines", []) or []
            default_text = "\n".join([line.content for line in azure_lines])

            dictionary = load_json(DICT_FILE)
            gpt_checked_text = default_text if skip_gpt else gpt_fix_text(default_text, dictionary)

            tab1, tab2, tab3, tab4 = st.tabs(["📄 元ファイル", "🖨️ OCRテキスト", "🤖 GPT補正", "✍️ 手作業修正"])
            with tab1:
                st.image(clean_img, caption=f"元ファイル (ページ {page_num})", use_container_width=True)
            with tab2:
                st.text_area(f"OCRテキスト（ページ {page_num}）", default_text, height=320, key=f"ocr_{page_num}")
            with tab3:
                st.text_area(f"GPT補正（ページ {page_num}）", gpt_checked_text, height=320, key=f"gpt_{page_num}")
            with tab4:
                corrected_text = st.text_area(f"手作業修正（ページ {page_num}）", gpt_checked_text, height=320, key=f"edit_{page_num}")
                if st.button(f"修正を保存 (ページ {page_num})", key=f"save_{page_num}"):
                    learned = learn_charwise_with_missing(default_text, corrected_text)
                    if learned:
                        update_dictionary_and_untrained(learned)
                        st.success(f"辞書と学習候補に {len(learned)} 件を追加しました！")
                    else:
                        st.info("修正が検出されませんでした。")
                    st.rerun()

            final_text_page = (corrected_text or gpt_checked_text).strip()
            all_corrected_texts.append(final_text_page)

            gpt_lines = [ln for ln in (corrected_text or gpt_checked_text).splitlines()]
            lines_for_layout = []
            for i, ln in enumerate(azure_lines):
                x_min, x_max, y_min, y_max = line_bbox(ln)
                text_for_line = gpt_lines[i] if i < len(gpt_lines) else ln.content
                lines_for_layout.append({"text": text_for_line, "x_min": x_min, "x_max": x_max, "y": y_min})

            page_layout_info = {
                "page_width": getattr(doc_page, "width", None) or 1.0,
                "page_height": getattr(doc_page, "height", None) or 1.0,
                "unit": getattr(doc_page, "unit", None) or "pixel",
                "lines": lines_for_layout
            }
            pages_layout.append(page_layout_info)

            done += 1; progress.progress(done / total_to_process)

        del pages, page_numbers
        gc.collect()

    status.success("✅ すべてのページの処理が完了しました。")

    if all_corrected_texts:
        joined_txt = "\n\n".join(all_corrected_texts)
        st.download_button("📥 補正テキスト（TXT, 見出しなし）", data=joined_txt.encode("utf-8"),
                           file_name="ocr_corrected.txt", mime="text/plain")

    if pages_layout:
        try:
            docx_bytes, indent_summary = build_docx_from_layout(
                pages_layout,
                add_label=show_labels,
                center_thresh_ratio=center_pct / 100.0,
                indent_epsilon_cm=float(indent_eps)
            )
            st.download_button("📥 Word（.docx：整列/段を反映）", data=docx_bytes,
                               file_name="ocr_layout.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            if indent_summary:
                st.markdown("#### 検出されたインデント段")
                for s in indent_summary:
                    st.write(f"- 段 {s['index']} : 中心 {s['center_cm']} cm（{s['count']} 行）")
        except Exception as e:
            st.warning(f"Word出力に失敗しました：{e}")

else:
    # ==== 画像ファイル：1ページ処理 ====
    try:
        try:
            img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        except Exception:
            arr = np.frombuffer(file_bytes, dtype=np.uint8)
            bgr = cv2.imdecode(arr, cv2.IMREAD_COLOR)
            if bgr is None:
                raise ValueError("画像の読み込みに失敗しました（JPG/PNG/PDFのみ対応）。")
            img = Image.fromarray(cv2.cvtColor(bgr, cv2.COLOR_BGR2RGB))
        pages = [img]; page_numbers = [1]
    except Exception as e:
        st.exception(e); st.stop()

    center_pct = 6; indent_eps = 0.6; show_labels = False
    all_corrected_texts: List[str] = []; pages_layout: List[Dict[str, Any]] = []

    for page_img, page_num in zip(pages, page_numbers):
        st.write(f"## ページ {page_num}")
        clean_img = remove_red_stamp(page_img)

        buf = io.BytesIO(); clean_img.save(buf, format="PNG"); buf.seek(0)

        with st.spinner("OCRを実行中..."):
            t0 = time.perf_counter()
            try:
                poller = client.begin_analyze_document("prebuilt-read", document=buf)
                result = poller.result(timeout=float(ocr_timeout))
            except Exception as e:
                st.error(f"OCRが{ocr_timeout}秒以内に完了しませんでした / 失敗しました：{e}")
                st.caption(f"OCR実行時間: {time.perf_counter() - t0:.1f}s")
                continue
            st.caption(f"OCR実行時間: {time.perf_counter() - t0:.1f}s")

        doc_page = result.pages[0] if getattr(result, "pages", None) else None
        if not doc_page:
            st.warning("OCR結果にページが見つかりませんでした。"); continue

        azure_lines = getattr(doc_page, "lines", []) or []
        default_text = "\n".join([line.content for line in azure_lines])

        dictionary = load_json(DICT_FILE)
        gpt_checked_text = default_text if skip_gpt else gpt_fix_text(default_text, dictionary)

        tab1, tab2, tab3, tab4 = st.tabs(["📄 元ファイル", "🖨️ OCRテキスト", "🤖 GPT補正", "✍️ 手作業修正"])
        with tab1:
            st.image(clean_img, caption=f"元ファイル (ページ {page_num})", use_container_width=True)
        with tab2:
            st.text_area(f"OCRテキスト（ページ {page_num}）", default_text, height=320, key=f"ocr_{page_num}")
        with tab3:
            st.text_area(f"GPT補正（ページ {page_num}）", gpt_checked_text, height=320, key=f"gpt_{page_num}")
        with tab4:
            corrected_text = st.text_area(f"手作業修正（ページ {page_num}）", gpt_checked_text, height=320, key=f"edit_{page_num}")
            if st.button(f"修正を保存 (ページ {page_num})", key=f"save_{page_num}"):
                learned = learn_charwise_with_missing(default_text, corrected_text)
                if learned:
                    update_dictionary_and_untrained(learned)
                    st.success(f"辞書と学習候補に {len(learned)} 件を追加しました！")
                else:
                    st.info("修正が検出されませんでした。")
                st.rerun()

        final_text_page = (corrected_text or gpt_checked_text).strip()
        all_corrected_texts.append(final_text_page)

        gpt_lines = [ln for ln in (corrected_text or gpt_checked_text).splitlines()]
        lines_for_layout = []
        for i, ln in enumerate(azure_lines):
            x_min, x_max, y_min, y_max = line_bbox(ln)
            text_for_line = gpt_lines[i] if i < len(gpt_lines) else ln.content
            lines_for_layout.append({"text": text_for_line, "x_min": x_min, "x_max": x_max, "y": y_min})

        page_layout_info = {
            "page_width": getattr(doc_page, "width", None) or 1.0,
            "page_height": getattr(doc_page, "height", None) or 1.0,
            "unit": getattr(doc_page, "unit", None) or "pixel",
            "lines": lines_for_layout
        }
        pages_layout.append(page_layout_info)

    if all_corrected_texts:
        joined_txt = "\n\n".join(all_corrected_texts)
        st.download_button("📥 補正テキスト（TXT, 見出しなし）", data=joined_txt.encode("utf-8"),
                           file_name="ocr_corrected.txt", mime="text/plain")

    if pages_layout:
        try:
            docx_bytes, indent_summary = build_docx_from_layout(
                pages_layout, add_label=show_labels, center_thresh_ratio=center_pct/100.0, indent_epsilon_cm=float(indent_eps)
            )
            st.download_button("📥 Word（.docx：整列/段を反映）", data=docx_bytes,
                               file_name="ocr_layout.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            if indent_summary:
                st.markdown("#### 検出されたインデント段")
                for s in indent_summary:
                    st.write(f"- 段 {s['index']} : 中心 {s['center_cm']} cm（{s['count']} 行）")
        except Exception as e:
            st.warning(f"Word出力に失敗しました：{e}")
