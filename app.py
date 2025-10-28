# dummy line
import streamlit as st
import os
import io
import json
import time
import gc
import re
import hashlib
import unicodedata
from itertools import zip_longest
from typing import List, Tuple, Dict, Any, Optional

import numpy as np
import cv2
from PIL import Image, ImageFile
import pypdfium2 as pdfium

from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential

# OpenAI（APIキー未設定でも起動できるよう防御）
try:
    from openai import OpenAI
except Exception:
    OpenAI = None

# PyTorch（任意・存在すれば使う）
try:
    import torch
    import torch.nn as nn
    TORCH_OK = True
except Exception:
    torch = None
    nn = None
    TORCH_OK = False

from docx import Document
from docx.shared import Cm, Pt

# ===================== 基本設定 =====================
AZURE_DOCINT_ENDPOINT = os.getenv("AZURE_DOCINT_ENDPOINT") or st.secrets.get("AZURE_DOCINT_ENDPOINT")
AZURE_DOCINT_KEY = os.getenv("AZURE_DOCINT_KEY") or st.secrets.get("AZURE_DOCINT_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY")
MODEL_NAME = os.getenv("OCR_GPT_MODEL") or st.secrets.get("OCR_GPT_MODEL", "gpt-5")
BATCH_SIZE_DEFAULT = max(1, int(os.getenv("OCR_BATCH_PAGES") or st.secrets.get("OCR_BATCH_PAGES", "10")))

if not AZURE_DOCINT_ENDPOINT or not AZURE_DOCINT_KEY:
    st.error("環境変数またはSecretsで AZURE_DOCINT_ENDPOINT / AZURE_DOCINT_KEY を設定してください。")
    st.stop()

client = DocumentAnalysisClient(endpoint=AZURE_DOCINT_ENDPOINT, credential=AzureKeyCredential(AZURE_DOCINT_KEY))

if OPENAI_API_KEY and OpenAI is not None:
    openai_client = OpenAI(api_key=OPENAI_API_KEY)
else:
    openai_client = None  # 未設定でもアプリは動く（GPT補正をスキップ）

JP_CHAR_RE = re.compile(r"^[\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FFF]$")

# ===================== ストレージ選択（local / azureblob） =====================
STORAGE_BACKEND = os.getenv("OCR_DICT_BACKEND") or st.secrets.get("OCR_DICT_BACKEND", "local")

if STORAGE_BACKEND == "azureblob":
    try:
        from azure.storage.blob import BlobServiceClient, ContentSettings
        from azure.core.exceptions import ResourceExistsError
    except Exception as e:
        st.error(f"azure-storage-blob が必要です。requirements.txt に追加してください。詳細: {e}")
        st.stop()

    AZURE_BLOB_CONN_STR = os.getenv("AZURE_BLOB_CONN_STR") or st.secrets.get("AZURE_BLOB_CONN_STR")
    OCR_DICT_CONTAINER = os.getenv("OCR_DICT_CONTAINER") or st.secrets.get("OCR_DICT_CONTAINER", "ocr-shared-dict")
    if not AZURE_BLOB_CONN_STR:
        st.error("Secrets/環境変数 AZURE_BLOB_CONN_STR を設定してください（Blob接続文字列）。")
        st.stop()

    _blob_service = BlobServiceClient.from_connection_string(AZURE_BLOB_CONN_STR)
    _container = _blob_service.get_container_client(OCR_DICT_CONTAINER)
    try:
        _container.create_container()
    except ResourceExistsError:
        pass

    # 辞書JSONの論理キー（Blob名）
    DICT_FILE = "ocr_char_corrections.json"
    UNTRAINED_FILE = "untrained_confusions.json"
    TRAINED_FILE = "trained_confusions.json"
    EVIDENCE_DIR = "evidence"  # 画像・メタ保存用のプレフィックス

    def _load_json_blob(blob_name: str) -> dict:
        try:
            data = _container.download_blob(blob_name).readall()
            return json.loads(data.decode("utf-8"))
        except Exception:
            return {}

    def _save_json_blob(obj: dict, blob_name: str):
        b = json.dumps(obj, ensure_ascii=False, indent=2).encode("utf-8")
        _container.upload_blob(
            name=blob_name, data=b, overwrite=True,
            content_settings=ContentSettings(content_type="application/json; charset=utf-8"),
        )

    def _upload_bytes(path_name: str, b: bytes, mime: str):
        _container.upload_blob(
            name=path_name, data=b, overwrite=True,
            content_settings=ContentSettings(content_type=mime),
        )

    def load_json_any(key: str) -> dict:
        return _load_json_blob(key)

    def save_json_any(obj: dict, key: str):
        _save_json_blob(obj, key)

    def save_bytes_any(path_name: str, b: bytes, mime: str = "application/octet-stream"):
        _upload_bytes(path_name, b, mime)

elif STORAGE_BACKEND == "local":
    DICT_DIR = os.getenv("OCR_DICT_DIR") or st.secrets.get("OCR_DICT_DIR", ".")
    DICT_FILE = os.path.join(DICT_DIR, "ocr_char_corrections.json")
    UNTRAINED_FILE = os.path.join(DICT_DIR, "untrained_confusions.json")
    TRAINED_FILE = os.path.join(DICT_DIR, "trained_confusions.json")
    EVIDENCE_DIR = os.path.join(DICT_DIR, "evidence")

    def _load_json_local(path: str, retries: int = 3, delay: float = 0.1) -> dict:
        for _ in range(retries):
            try:
                with open(path, "r", encoding="utf-8") as f:
                    return json.load(f)
            except FileNotFoundError:
                return {}
            except (json.JSONDecodeError, PermissionError):
                time.sleep(delay)
        return {}

    def _save_json_local(obj: dict, path: str):
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        tmp_path = f"{path}.tmp"
        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(obj, f, ensure_ascii=False, indent=2)
            f.flush(); os.fsync(f.fileno())
        os.replace(tmp_path, path)

    def _save_bytes_local(path: str, b: bytes):
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        with open(path, "wb") as f:
            f.write(b); f.flush(); os.fsync(f.fileno())

    def load_json_any(key: str) -> dict:
        return _load_json_local(key)

    def save_json_any(obj: dict, key: str):
        _save_json_local(obj, key)

    def save_bytes_any(path_name: str, b: bytes, mime: str = "application/octet-stream"):
        # mime はダミー。ローカルでは単に保存
        _save_bytes_local(path_name, b)
else:
    st.error(f"未知の OCR_DICT_BACKEND: {STORAGE_BACKEND}（local / azureblob のみ対応）")
    st.stop()

# ===================== ユーティリティ =====================
def remove_red_stamp(img_pil: Image.Image) -> Image.Image:
    """シンプル版：赤系を白で抑制（印影除去ON時のみ使用）"""
    img = np.array(img_pil)
    hsv = cv2.cvtColor(img, cv2.COLOR_RGB2HSV)
    lower_red1 = np.array([0, 70, 50]); upper_red1 = np.array([10, 255, 255])
    lower_red2 = np.array([170, 70, 50]); upper_red2 = np.array([180, 255, 255])
    mask = cv2.bitwise_or(cv2.inRange(hsv, lower_red1, upper_red1), cv2.inRange(hsv, lower_red2, upper_red2))
    img[mask > 0] = [255, 255, 255]
    return Image.fromarray(img)

def learn_charwise_with_missing(original: str, corrected: str) -> dict:
    """文字単位で original→corrected の差分を抽出し、{誤: {right: 正, count}} を返す"""
    learned: Dict[str, Dict[str, Any]] = {}
    import difflib
    sm = difflib.SequenceMatcher(None, original, corrected)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag in ["replace", "insert"]:
            o_seg, c_seg = original[i1:i2], corrected[j1:j2]
            for o_char, c_char in zip_longest(o_seg, c_seg, fillvalue=""):
                if c_char and (not o_char or o_char != c_char):
                    wrong = o_char if o_char else "□"
                    if JP_CHAR_RE.match(c_char) or c_char == "□":
                        learned[wrong] = {"right": c_char, "count": 1}
    return learned

def update_dictionary_and_untrained(learned: dict):
    """辞書・学習候補JSONを更新"""
    dictionary = load_json_any(DICT_FILE)
    for w, meta in learned.items():
        if w in dictionary:
            if dictionary[w]["right"] == meta["right"]:
                dictionary[w]["count"] += meta["count"]
            else:
                if meta["count"] > dictionary[w]["count"]:
                    dictionary[w] = meta
        else:
            dictionary[w] = meta
    save_json_any(dictionary, DICT_FILE)

    untrained = load_json_any(UNTRAINED_FILE)
    for w, meta in learned.items():
        untrained[w] = meta["right"]
    save_json_any(untrained, UNTRAINED_FILE)

def gpt_fix_text(text: str, dictionary: dict) -> str:
    """GPTで最小修正（辞書をプロンプトに同梱）"""
    if openai_client is None:
        st.info("OPENAI_API_KEY が未設定のため、GPT補正はスキップします。")
        return text

    prompt = f"""
次のOCR結果を自然な日本語に直してください。
- 日本語に存在しない文字は「□」にしてください。
- 辞書候補を参考にしてください: {json.dumps(dictionary, ensure_ascii=False)}
- 意味を勝手に補完せず、最小限の修正だけ行ってください。

OCR結果:
{text}
""".strip()
    try:
        resp = openai_client.responses.create(
            model=MODEL_NAME,
            input=prompt,
            text={"verbosity": "low"},
            reasoning={"effort": "minimal"},
        )
        if hasattr(resp, "output_text") and resp.output_text:
            return resp.output_text.strip()
        out_parts = []
        for item in getattr(resp, "output", []) or []:
            for part in getattr(item, "content", []) or []:
                t = getattr(part, "text", None)
                if t: out_parts.append(t)
        out = "".join(out_parts).strip()
        return out or text
    except Exception as e:
        st.warning(f"GPT補正をスキップしました（エラー）：{e}")
        return text

# ====== ページ指定の正規化・解析（全角対応） ======
def parse_page_spec(spec: str, max_pages: int) -> List[int]:
    """
    '1,3,5-7' を 0始まりインデックスへ。
    全角数字/カンマ/各種ダッシュも許可（'１，３，５－７' '1—3' など）
    """
    s = (spec or "").strip()
    if not s:
        return []
    s = unicodedata.normalize("NFKC", s)
    s = s.replace("，", ",").replace("、", ",")
    for dash in ["－", "ー", "―", "—", "–"]:
        s = s.replace(dash, "-")
    out = set()
    parts = [p.strip() for p in s.split(",") if p.strip()]
    for p in parts:
        if "-" in p:
            a, b = p.split("-", 1)
            try:
                ia = int(a); ib = int(b)
                lo, hi = min(ia, ib), max(ia, ib)
                lo = max(1, lo); hi = min(max_pages, hi)
                for n in range(lo, hi + 1):
                    out.add(n - 1)
            except ValueError:
                continue
        else:
            try:
                n = int(p)
                if 1 <= n <= max_pages:
                    out.add(n - 1)
            except ValueError:
                continue
    return sorted(out)

def chunked(seq: List[int], n: int) -> List[List[int]]:
    return [seq[i:i+n] for i in range(0, len(seq), n)]

# ========== Azure行の座標・bbox ==========
def _line_bbox(line_obj: Any) -> Tuple[float, float, float, float]:
    poly = getattr(line_obj, "polygon", None) or getattr(line_obj, "bounding_polygon", None)
    if not poly:
        # 古いAPIでは bounding_regions から取る場合もあるが簡略化
        return (0.0, 0.0, 0.0, 0.0)
    xs, ys = [], []
    for p in poly:
        x = getattr(p, "x", None); y = getattr(p, "y", None)
        if x is None and isinstance(p, dict):
            x = p.get("x", 0.0); y = p.get("y", 0.0)
        xs.append(float(x)); ys.append(float(y))
    return (min(xs or [0.0]), min(ys or [0.0]), max(xs or [0.0]), max(ys or [0.0]))

# ========== PDFレンダリング ==========
def render_pdf_selected_pages(pdf_bytes: bytes, indices_0based: List[int], dpi: int = 200) -> Tuple[List[Image.Image], List[int]]:
    imgs: List[Image.Image] = []
    nums: List[int] = []
    pdf = pdfium.PdfDocument(io.BytesIO(pdf_bytes))
    scale = dpi / 72.0
    for idx in indices_0based:
        page = pdf[idx]
        pil = page.render(scale=scale).to_pil().convert("RGB")
        imgs.append(pil); nums.append(idx + 1)
    return imgs, nums

# ========== OCRコア（Azure polling） ==========
def _ocr_polling(png_bytes: bytes, timeout_sec: float) -> dict:
    poller = client.begin_analyze_document("prebuilt-read", document=io.BytesIO(png_bytes))
    t0 = time.perf_counter()
    status = st.empty()
    while True:
        if poller.done():
            break
        elapsed = time.perf_counter() - t0
        if elapsed > float(timeout_sec):
            status.empty()
            raise TimeoutError(f"Azure OCR timeout after {elapsed:.1f}s")
        status.info(f"Azure OCR 実行中… {elapsed:.1f}s / {timeout_sec:.0f}s")
        time.sleep(0.3)
    status.empty()

    result = poller.result()
    doc_page = result.pages[0] if getattr(result, "pages", None) else None
    if not doc_page:
        return {"pw": 1.0, "ph": 1.0, "lines": [], "raw": getattr(result, "content", "")}

    lines = []
    for ln in getattr(doc_page, "lines", []) or []:
        x1, y1, x2, y2 = _line_bbox(ln)
        lines.append({
            "content": ln.content,
            "x1": float(x1), "y1": float(y1), "x2": float(x2), "y2": float(y2),
            "x": float(x1), "y": float(y1)  # 後方互換
        })

    return {
        "pw": float(getattr(doc_page, "width", 1.0) or 1.0),
        "ph": float(getattr(doc_page, "height", 1.0) or 1.0),
        "lines": lines,
        "raw": getattr(result, "content", "") or "\n".join([l["content"] for l in lines]),
    }

@st.cache_data(show_spinner=False)
def _ocr_cached(digest: str, png_bytes: bytes, timeout_sec: float) -> dict:
    return _ocr_polling(png_bytes, timeout_sec)

def _ocr_dispatch(png_bytes: bytes, timeout_sec: float, use_cache: bool) -> dict:
    if use_cache:
        digest = hashlib.md5(png_bytes).hexdigest()
        return _ocr_cached(digest, png_bytes, timeout_sec)
    else:
        return _ocr_polling(png_bytes, timeout_sec)

# ===================== 証拠ログ：行画像＆文字パッチの保存 =====================
def _bytes_from_pil(img: Image.Image, fmt="PNG") -> bytes:
    b = io.BytesIO()
    img.save(b, format=fmt)
    return b.getvalue()

def _safe_int(x: float) -> int:
    try:
        return int(round(float(x)))
    except Exception:
        return 0

def _crop_safe(img: Image.Image, box: Tuple[int,int,int,int]) -> Image.Image:
    W, H = img.size
    x1, y1, x2, y2 = box
    x1 = max(0, min(W, x1)); x2 = max(0, min(W, x2))
    y1 = max(0, min(H, y1)); y2 = max(0, min(H, y2))
    if x2 <= x1 or y2 <= y1:
        return img.crop((0,0,1,1))
    return img.crop((x1, y1, x2, y2))

def _estimate_char_box(line: dict, idx_in_line: int, expand_px: int = 2) -> Tuple[int,int,int,int]:
    """行bboxと行テキスト長から、対象文字の近似bboxを推定（横書き前提・等幅近似）"""
    x1, y1, x2, y2 = line.get("x1", 0), line.get("y1", 0), line.get("x2", 0), line.get("y2", 0)
    L = max(1, len(line.get("content", "")))
    w = max(1, x2 - x1)
    ch_w = w / L
    cx1 = int(x1 + idx_in_line * ch_w) - expand_px
    cx2 = int(x1 + (idx_in_line + 1) * ch_w) + expand_px
    cy1 = int(y1) - expand_px
    cy2 = int(y2) + expand_px
    return (cx1, cy1, cx2, cy2)

def _ensure_date_prefix() -> str:
    from datetime import datetime
    return datetime.now().strftime("%Y%m%d")

def save_evidence_assets(
    backend_prefix: str,
    page_num: int,
    page_png_bytes_raw: bytes,
    azure_lines: List[dict],
    ocr_text: str,
    corrected_text: str,
    evidence_on: bool = True,
) -> Tuple[int, Optional[str]]:
    """
    学習用に、行画像と差分の近似文字パッチを保存。
    返り値: (保存したファイル点数, エラー文字列)
    """
    if not evidence_on:
        return (0, None)
    try:
        date_prefix = _ensure_date_prefix()
        # ルート（Blob名 or ローカルパス）
        root = f"{EVIDENCE_DIR}/{date_prefix}/page-{page_num:03d}"
        # ページ原画
        save_bytes_any(f"{root}/page.png", page_png_bytes_raw, "image/png")

        # 行ごとの画像
        img = Image.open(io.BytesIO(page_png_bytes_raw)).convert("RGB")
        saved = 1  # page.png

        # 差分抽出（行単位）
        import difflib
        # 行をやや素朴に分割（Azureは行リストを持つため、それを優先）
        # ただし補助として ocr_text 全体も保存
        save_bytes_any(f"{root}/ocr_text.txt", ocr_text.encode("utf-8"), "text/plain")
        save_bytes_any(f"{root}/corrected_text.txt", corrected_text.encode("utf-8"), "text/plain")
        saved += 2

        for li, line in enumerate(azure_lines):
            txt = line.get("content", "")
            x1, y1, x2, y2 = map(_safe_int, (line.get("x1", 0), line.get("y1", 0), line.get("x2", 0), line.get("y2", 0)))
            line_img = _crop_safe(img, (x1, y1, x2, y2))
            save_bytes_any(f"{root}/lines/line_{li:03d}.png", _bytes_from_pil(line_img), "image/png")
            save_bytes_any(f"{root}/lines/line_{li:03d}.txt", txt.encode("utf-8"), "text/plain")
            saved += 2

        # 文字パッチ：original vs corrected の差分を、近似座標で切り出し
        # 行対応：Azure行配列順のまま、正規化して突き合わせ（簡易）
        # ここでは corrected_text を行数で丸める簡易実装
        corrected_lines = corrected_text.splitlines()
        for li, line in enumerate(azure_lines):
            if li >= len(corrected_lines):
                break
            orig_line = line.get("content", "")
            corr_line = corrected_lines[li]
            sm = difflib.SequenceMatcher(None, orig_line, corr_line)
            for tag, i1, i2, j1, j2 in sm.get_opcodes():
                if tag in ["replace", "insert"]:
                    for k, (o_char, c_char) in enumerate(zip_longest(orig_line[i1:i2], corr_line[j1:j2], fillvalue="")):
                        if c_char and (not o_char or o_char != c_char):
                            # 近似位置でパッチ切り出し
                            char_idx = i1 + k
                            cx1, cy1, cx2, cy2 = _estimate_char_box(line, char_idx, expand_px=2)
                            patch = _crop_safe(img, (cx1, cy1, cx2, cy2))
                            save_bytes_any(
                                f"{root}/patches/line_{li:03d}_char_{char_idx:03d}_{o_char or '□'}->{c_char}.png",
                                _bytes_from_pil(patch),
                                "image/png"
                            )
                            saved += 1

        # メタ
        meta = {
            "backend": STORAGE_BACKEND,
            "page": page_num,
            "lines": [{"text": L.get("content",""), "bbox":[L.get("x1",0),L.get("y1",0),L.get("x2",0),L.get("y2",0)]} for L in azure_lines]
        }
        save_bytes_any(f"{root}/meta.json", json.dumps(meta, ensure_ascii=False, indent=2).encode("utf-8"), "application/json")
        saved += 1

        return (saved, None)
    except Exception as e:
        return (0, str(e))

# ===================== 文字分類器（任意）：読み込み＆推論 =====================
@st.cache_resource(show_spinner=False)
def _load_char_classifier() -> Optional[Any]:
    """
    文字分類器（例: ResNet18）を任意読み込み。存在しない/依存がない場合は None。
    期待する重みファイル名は環境変数または既定 'char_classifier_resnet18.pt'
    """
    if not TORCH_OK:
        return None
    model_path = os.getenv("CHAR_CLS_PATH") or st.secrets.get("CHAR_CLS_PATH", "char_classifier_resnet18.pt")
    if not os.path.exists(model_path):
        return None
    try:
        # 単純な畳み込みネット（クラス数や入出力形状は重みに依存するため、TorchScript推奨）
        # ここでは state_dict をロードできる前提でラッパのみ用意（不一致ならNoneで運用）
        model = torch.jit.load(model_path, map_location="cpu") if model_path.endswith(".ts") else None
        if model is None:
            # フォールバック：state_dict ロードに失敗しやすいので try/except
            # ここでは安全に None を返す
            return None
        model.eval()
        return model
    except Exception:
        return None

def _prep_patch_for_model(patch: Image.Image, size: int = 28) -> Optional["torch.Tensor"]:
    if not TORCH_OK:
        return None
    try:
        g = patch.convert("L").resize((size, size))
        arr = np.array(g).astype("float32") / 255.0
        t = torch.from_numpy(arr).unsqueeze(0).unsqueeze(0)  # [1,1,H,W]
        return t
    except Exception:
        return None

def _classify_patch(model, patch: Image.Image) -> Tuple[Optional[int], float]:
    """モデルでパッチを分類して top-1 と信頼度（softmax）を返す。失敗時は (None, 0.0)"""
    if (not TORCH_OK) or (model is None):
        return (None, 0.0)
    try:
        t = _prep_patch_for_model(patch)
        if t is None:
            return (None, 0.0)
        with torch.no_grad():
            logits = model(t)  # 期待: [1, C]
            if not isinstance(logits, torch.Tensor) or logits.ndim != 2:
                return (None, 0.0)
            probs = torch.softmax(logits, dim=1).cpu().numpy()[0]
            pred_id = int(np.argmax(probs))
            conf = float(np.max(probs))
            return pred_id, conf
    except Exception:
        return (None, 0.0)

# 文字ID→Unicodeのマップ（本来は学習時に保存しておく）
# ない場合は分類器をアクティブにしない
def _load_labelmap() -> Optional[Dict[int, str]]:
    # 簡易：secrets/jsonファイルなどから読み込み可能に
    try:
        path = os.getenv("CHAR_LABELMAP_JSON") or st.secrets.get("CHAR_LABELMAP_JSON", "")
        if path and os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                d = json.load(f)
            # キーをintに
            return {int(k): v for k, v in d.items()}
    except Exception:
        pass
    return None

# 自動置換（分類器）本体
def auto_replace_with_classifier(
    page_img_raw: Image.Image,
    azure_lines: List[dict],
    base_text: str,
    confusion_dict: Dict[str, str],
    cls_model,
    labelmap: Optional[Dict[int, str]],
    conf_thresh: float = 0.9
) -> Tuple[str, List[Dict[str, Any]]]:
    """
    base_text（ページ全体のテキスト）に対して、
    混同文字（confusion_dictキー）だけ近似bboxでパッチを切って分類し、高信頼なら置換。
    返り値: (置換後テキスト, 置換リスト[{pos,char_from,char_to,conf}]）
    """
    if (not TORCH_OK) or (cls_model is None) or (labelmap is None) or not confusion_dict:
        return base_text, []

    # 行単位で処理
    new_lines: List[str] = []
    repl_log: List[Dict[str, Any]] = []
    base_lines = base_text.splitlines()
    for li, line in enumerate(azure_lines):
        src = base_lines[li] if li < len(base_lines) else line.get("content", "")
        if not src:
            new_lines.append(src)
            continue
        x1, y1, x2, y2 = map(_safe_int, (line.get("x1",0), line.get("y1",0), line.get("x2",0), line.get("y2",0)))
        # 行画像
        # ただしパッチはページ原画から切り出す（等幅近似なのでページからでも変わらない）
        row = list(src)
        changed = False
        for idx, ch in enumerate(row):
            if ch in confusion_dict:
                # 近似で文字位置を切り出し
                cx1, cy1, cx2, cy2 = _estimate_char_box(line, idx, expand_px=2)
                patch = _crop_safe(page_img_raw, (cx1, cy1, cx2, cy2))
                pred_id, conf = _classify_patch(cls_model, patch)
                if pred_id is None or conf < conf_thresh:
                    continue
                pred_char = labelmap.get(pred_id)
                if not pred_char:
                    continue
                desired = confusion_dict[ch]  # 推奨の正字
                # 分類器の予測が推奨正字と一致し、かつ高信頼なら置換
                if pred_char == desired:
                    row[idx] = desired
                    changed = True
                    repl_log.append({"line": li, "index": idx, "from": ch, "to": desired, "conf": conf})
        new_lines.append("".join(row) if changed else src)

    return "\n".join(new_lines), repl_log

# ===================== Word生成 =====================
EMU_PER_CM = 360000.0
def to_cm(val) -> float:
    try:
        return float(getattr(val, "cm"))
    except Exception:
        try:
            return float(val) / EMU_PER_CM
        except Exception:
            return float(val)

def build_docx_from_layout(pages_layout: List[Dict[str, Any]]) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0); section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)

    page_w_cm = to_cm(section.page_width)
    left_cm = to_cm(section.left_margin); right_cm = to_cm(section.right_margin)
    content_width_cm = max(0.1, page_w_cm - left_cm - right_cm)

    style = doc.styles["Normal"]; style.font.name = "Yu Gothic"; style.font.size = Pt(11)

    for idx, page in enumerate(pages_layout, start=1):
        pw = float(page.get("page_width") or 1.0)
        ph = float(page.get("page_height") or 1.0)
        lines = page.get("lines", [])
        lines_sorted = sorted(lines, key=lambda r: (r["y"], r["x"]))
        y_thresh = ph * 0.018
        prev_y = None

        for item in lines_sorted:
            txt = item["text"]; x = float(item["x"]); y = float(item["y"])
            para = doc.add_paragraph(); para.add_run(txt)
            indent_cm = max(0.0, min(0.9 * content_width_cm, (x / max(pw, 1e-6)) * content_width_cm))
            para.paragraph_format.left_indent = Cm(indent_cm)
            para.paragraph_format.space_after = Pt(2)
            if prev_y is not None and (y - prev_y) > y_thresh:
                para.paragraph_format.space_before = Pt(8)
            prev_y = y

        if idx < len(pages_layout):
            doc.add_page_break()

    bio = io.BytesIO(); doc.save(bio); bio.seek(0)
    return bio.read()

# ===================== UI =====================
st.title("📄 Document Intelligence OCR（Azure）— 学習ログ/分類器/印影ON-OFF/ページ指定/バッチ/GPT/Word")

# 診断
st.sidebar.markdown("### 🔧 環境")
st.sidebar.write({
    "AZURE_DOCINT_ENDPOINT_set": bool(AZURE_DOCINT_ENDPOINT),
    "AZURE_DOCINT_KEY_set": bool(AZURE_DOCINT_KEY),
    "OPENAI_API_KEY_set": bool(OPENAI_API_KEY),
    "MODEL": MODEL_NAME,
    "BATCH_SIZE(default)": BATCH_SIZE_DEFAULT,
})

st.sidebar.markdown("### 📂 辞書の参照先")
if STORAGE_BACKEND == "azureblob":
    st.sidebar.write({
        "OCR_DICT_BACKEND": "azureblob",
        "container": os.getenv("OCR_DICT_CONTAINER") or st.secrets.get("OCR_DICT_CONTAINER", "ocr-shared-dict"),
        "DICT_BLOB": DICT_FILE, "UNTRAINED_BLOB": UNTRAINED_FILE, "TRAINED_BLOB": TRAINED_FILE,
        "EVIDENCE_DIR": EVIDENCE_DIR,
    })
else:
    st.sidebar.write({
        "OCR_DICT_BACKEND": "local",
        "OCR_DICT_DIR": os.path.abspath(os.getenv("OCR_DICT_DIR") or st.secrets.get("OCR_DICT_DIR", ".")),
        "DICT_FILE": os.path.abspath(DICT_FILE),
        "UNTRAINED_FILE": os.path.abspath(UNTRAINED_FILE),
        "TRAINED_FILE": os.path.abspath(TRAINED_FILE),
        "EVIDENCE_DIR": os.path.abspath(EVIDENCE_DIR),
    })

# デバッグUI
st.sidebar.markdown("### 🛠 デバッグ")
skip_gpt = st.sidebar.checkbox("GPT補正をスキップ", value=(openai_client is None))
ocr_timeout = st.sidebar.slider("OCRタイムアウト（秒）", 10, 120, 45, step=5)
batch_size_override = st.sidebar.number_input("バッチサイズ上書き", 1, 20, value=BATCH_SIZE_DEFAULT)
use_cache = st.sidebar.checkbox("OCRキャッシュ（実験）", value=False)
debug_log = st.sidebar.checkbox("🔍 詳細ログ", value=True)

# 分類器オプション
st.sidebar.markdown("### 🤖 文字分類器（任意）")
classifier_enable = st.sidebar.checkbox("🧪 混同文字の自動置換（高信頼のみ）", value=False)
conf_thresh = st.sidebar.slider("信頼度しきい値", 0.50, 0.99, 0.90, step=0.01)
if classifier_enable:
    cls_model = _load_char_classifier()
    labelmap = _load_labelmap()
else:
    cls_model = None
    labelmap = None
st.sidebar.write({"classifier_loaded": bool(cls_model), "labelmap_loaded": bool(labelmap)})

# 学習ログオプション
st.sidebar.markdown("### 📝 学習ログ（証拠の保存）")
evidence_on = st.sidebar.checkbox("修正保存時に 行画像＆文字パッチ を保存", value=True)

# 辞書プレビュー & 初期化
dict_preview_box = st.sidebar.container()
with dict_preview_box:
    st.subheader("📖 現在の辞書（プレビュー）")
    st.json(load_json_any(DICT_FILE))

def _init_dict_files():
    try:
        save_json_any({}, DICT_FILE)
        save_json_any({}, UNTRAINED_FILE)
        save_json_any({}, TRAINED_FILE)
        return True, None
    except Exception as e:
        return False, str(e)

colA, colB = st.sidebar.columns(2)
with colA:
    if st.button("🔄 辞書プレビュー再読込", type="secondary"):
        with dict_preview_box:
            st.subheader("📖 現在の辞書（プレビュー）")
            st.json(load_json_any(DICT_FILE))
with colB:
    if st.button("🧰 辞書ファイル初期化/作成"):
        ok, err = _init_dict_files()
        if ok:
            st.sidebar.success("辞書ファイルを初期化しました。『再読込』で確認してください。")
        else:
            st.sidebar.error(f"初期化に失敗: {err}")

# ===================== アップロード保持 =====================
if "file_bytes" not in st.session_state:
    st.session_state["file_bytes"] = None
    st.session_state["file_name"] = None
    st.session_state["file_mime"] = None
    st.session_state["is_pdf"] = False
    st.session_state["dpi"] = 200
    st.session_state["page_indices"] = []  # ← これの有無で実行済みか判断
    st.session_state["stamp_mode"] = "OFF"  # 印影除去モード

uploaded = st.file_uploader("画像またはPDFをアップロードしてください", type=["jpg", "jpeg", "png", "pdf"], key="uploader")
if uploaded is not None:
    st.session_state["file_bytes"] = uploaded.getvalue()
    st.session_state["file_name"] = uploaded.name
    st.session_state["file_mime"] = uploaded.type
    st.session_state["is_pdf"] = (uploaded.type == "application/pdf") or uploaded.name.lower().endswith(".pdf")
    # 新規アップロード時は、ページ選択・結果を初期化（印影モードは維持）
    for k in list(st.session_state.keys()):
        if k.startswith("ocr_") or k.startswith("gpt_") or k.startswith("edit_") or k.startswith("lines_") or k.startswith("imgraw_"):
            del st.session_state[k]
    st.session_state["page_indices"] = []

file_bytes = st.session_state["file_bytes"]
if not file_bytes:
    st.info("📂 ここにファイルをアップロードしてください")
    st.stop()

ImageFile.LOAD_TRUNCATED_IMAGES = True
is_input_pdf = st.session_state["is_pdf"]

# ====== サイドバー：実行ステータス（ran は page_indices の有無で算出）=====
ran_status = bool(st.session_state.get("page_indices"))
st.sidebar.markdown("### 📊 実行ステータス")
st.sidebar.write({
    "ran": ran_status,
    "has_file": bool(st.session_state.get("file_bytes")),
    "is_pdf": bool(is_input_pdf),
    "saved_indices": st.session_state.get("page_indices", []),
    "dpi": st.session_state.get("dpi", 200),
    "stamp_mode": st.session_state.get("stamp_mode", "OFF"),
})

# ===================== ページ選択（常時UI＋ボタンだが自動補完あり） =====================
if is_input_pdf:
    # PDFページ数
    try:
        pdf_for_count = pdfium.PdfDocument(io.BytesIO(file_bytes))
        total_pages = len(pdf_for_count)
        st.info(f"📘 このPDFは {total_pages} ページあります。")
    except Exception as e:
        st.exception(e); st.stop()

    st.subheader("▶ OCRするページを選択")
    col1, col2 = st.columns([2,1])
    with col1:
        select_mode = st.radio(
            "選択方法",
            options=["全ページ", "範囲指定", "ページ番号指定（例: 1,3,5-7 / １，３，５－７）"],
            index=1 if total_pages > 1 else 0,
            horizontal=True
        )
    with col2:
        dpi = st.slider("DPI", 72, 300, value=st.session_state.get("dpi", 200), step=4)
        st.session_state["dpi"] = dpi

    # --- 印影除去トグル（PDF用） ---
    stamp_toggle = st.radio(
        "印影除去（ハンコ）", ["OFF", "ON"], index=0, horizontal=True, key="ui_stamp_toggle"
    )

    if select_mode == "範囲指定" and total_pages > 1:
        start, end = st.slider("処理するページ範囲（1始まり）", 1, total_pages, (1, min(total_pages, 5)))
        chosen_indices = list(range(start - 1, end))
    elif select_mode == "ページ番号指定（例: 1,3,5-7 / １，３，５－７）":
        spec_default = "1-3" if total_pages >= 3 else "1"
        spec = st.text_input("ページ番号（カンマ区切り、範囲はハイフン。全角OK）", value=spec_default)
        chosen_indices = parse_page_spec(spec, total_pages)
    else:
        chosen_indices = list(range(total_pages))

    st.caption(f"選択中: {', '.join(str(i+1) for i in chosen_indices) if chosen_indices else '(なし)'} / DPI={dpi}")

    # 実行ボタン
    run_clicked = st.button("▶ この設定でOCRを実行", type="primary", key="run_pdf")

    # クリックされたら今回の選択と印影モードを保存
    if run_clicked:
        if chosen_indices:
            st.session_state["page_indices"] = chosen_indices
        st.session_state["stamp_mode"] = stamp_toggle

    # --- 自動補完：page_indices がまだ空なら埋める（押し忘れ・空入力対策） ---
    if not st.session_state.get("page_indices"):
        if chosen_indices:
            st.session_state["page_indices"] = chosen_indices
        else:
            st.session_state["page_indices"] = list(range(total_pages))  # 全ページで強制実行

    # 以降、この実行で使う設定を固定
    chosen_indices = st.session_state["page_indices"]
    stamp_mode = st.session_state.get("stamp_mode", stamp_toggle)  # 初回はUI値

    # ===== 実行本体 =====
    EFFECTIVE_BATCH = int(batch_size_override) if batch_size_override else BATCH_SIZE_DEFAULT
    total_to_process = len(chosen_indices)
    progress = st.progress(0.0)
    status_area = st.empty()
    all_corrected_texts: List[str] = []
    pages_layout: List[Dict[str, Any]] = []
    done = 0

    st.write("### ▶ 実行開始")
    st.write(f"🧪 ページ: {', '.join(str(i+1) for i in chosen_indices)} / DPI={dpi} / バッチ={EFFECTIVE_BATCH} / 印影: {stamp_mode}")

    for batch_no, batch_indices in enumerate(chunked(chosen_indices, EFFECTIVE_BATCH), start=1):
        status_area.info(f"🔄 バッチ {batch_no} / {((total_to_process - 1) // EFFECTIVE_BATCH) + 1} （ページ: {', '.join(str(i+1) for i in batch_indices)}）")
        try:
            pages, page_numbers = render_pdf_selected_pages(file_bytes, batch_indices, dpi=dpi)
        except Exception as e:
            st.exception(e); st.stop()

        for page_img, page_num in zip(pages, page_numbers):
            st.write(f"## ページ {page_num}")

            # 原画PNG（証拠用）を保存（印影除去せず）
            orig_png = _bytes_from_pil(page_img)

            # 印影ONなら赤抑制
            clean_img = remove_red_stamp(page_img) if stamp_mode == "ON" else page_img
            st.image(clean_img, caption=f"処理画像 (ページ {page_num})", use_container_width=True)

            # OCR画像は処理後のもの
            buf = io.BytesIO(); clean_img.save(buf, format="PNG"); png_bytes = buf.getvalue()

            with st.spinner("OCRを実行中..."):
                t0 = time.perf_counter()
                try:
                    cached = _ocr_dispatch(png_bytes, ocr_timeout, use_cache)
                except Exception as e:
                    st.error(f"OCRに失敗：{e}")
                    st.caption(f"OCR実行時間: {time.perf_counter() - t0:.1f}s")
                    continue
                elapsed = time.perf_counter() - t0
                st.caption(f"OCR実行時間: {elapsed:.1f}s")

            azure_lines = cached.get("lines") or []
            default_text = "\n".join([ln["content"] for ln in azure_lines]) if azure_lines else (cached.get("raw") or "")
            if not default_text.strip():
                st.warning("OCRは成功しましたがテキストが空でした。DPIや画像品質を見直してください。")

            # 分類器（任意）で混同文字を自動置換（高信頼のみ）
            dictionary = load_json_any(DICT_FILE)
            confusion_map = {k: v.get("right", v) if isinstance(v, dict) else v for k, v in dictionary.items()}
            auto_text = default_text
            repl_log = []
            if classifier_enable and cls_model is not None and labelmap is not None:
                auto_text, repl_log = auto_replace_with_classifier(
                    page_img_raw=page_img,  # 原画でパッチ切り出し
                    azure_lines=azure_lines,
                    base_text=default_text,
                    confusion_dict=confusion_map,
                    cls_model=cls_model,
                    labelmap=labelmap,
                    conf_thresh=conf_thresh
                )
                if repl_log:
                    st.info(f"🔁 自動置換 {len(repl_log)} 件（信頼度≥{conf_thresh:.2f}）")

            # GPT補正（スキップ可）：自動置換後のテキストを入力に
            base_for_gpt = auto_text
            gpt_checked_text = base_for_gpt if skip_gpt else gpt_fix_text(base_for_gpt, dictionary)

            # 状態キー
            ocr_key = f"ocr_{page_num}"
            gpt_key = f"gpt_{page_num}"
            edit_key = f"edit_{page_num}"
            lines_key = f"lines_{page_num}"
            imgraw_key = f"imgraw_{page_num}"

            # セッション保存（結果保持）
            if ocr_key not in st.session_state: st.session_state[ocr_key] = default_text
            if gpt_key not in st.session_state: st.session_state[gpt_key] = gpt_checked_text
            if edit_key not in st.session_state: st.session_state[edit_key] = gpt_checked_text
            st.session_state[lines_key] = azure_lines  # 後で証拠ログに使用
            st.session_state[imgraw_key] = orig_png    # 原画PNG

            # タブUI
            tab2, tab3, tab4 = st.tabs(["🖨️ OCRテキスト", "🤖 GPT/自動置換後", "✍️ 手作業修正"])
            with tab2:
                st.text_area(f"OCRテキスト（ページ {page_num}）", height=320, key=ocr_key)
                if repl_log:
                    st.caption(f"自動置換ログ: {repl_log[:3]}{' ...' if len(repl_log)>3 else ''}")
            with tab3:
                st.text_area(f"GPT/自動置換後（ページ {page_num}）", height=320, key=gpt_key)
            with tab4:
                st.text_area(f"手作業修正（ページ {page_num}）", height=320, key=edit_key)
                if st.button(f"修正を保存 (ページ {page_num})", key=f"save_{page_num}"):
                    corrected_text_current = st.session_state.get(edit_key, gpt_checked_text)
                    learned = learn_charwise_with_missing(st.session_state.get(ocr_key, default_text), corrected_text_current)
                    if learned:
                        update_dictionary_and_untrained(learned)
                        st.success(f"辞書と学習候補に {len(learned)} 件を追加しました！")
                    else:
                        st.info("修正が検出されませんでした。")

                    # 証拠ログの保存（行画像＋文字パッチ）
                    saved_cnt, err = save_evidence_assets(
                        backend_prefix=EVIDENCE_DIR,
                        page_num=page_num,
                        page_png_bytes_raw=st.session_state.get(imgraw_key, orig_png),
                        azure_lines=st.session_state.get(lines_key, azure_lines),
                        ocr_text=st.session_state.get(ocr_key, default_text),
                        corrected_text=corrected_text_current,
                        evidence_on=evidence_on
                    )
                    if err:
                        st.warning(f"学習証拠の保存に失敗：{err}")
                    else:
                        st.success(f"学習証拠を保存しました（{saved_cnt}ファイル）")

            # Wordレイアウト用に行座標＋テキストを蓄積（gpt/手修正を優先）
            final_text_page = (
                st.session_state.get(edit_key)
                or st.session_state.get(gpt_key)
                or gpt_checked_text
                or default_text
            ).strip()
            all_corrected_texts.append(final_text_page)

            gpt_lines = final_text_page.splitlines()
            lines_for_layout = []
            for i, ln in enumerate(azure_lines):
                x = float(ln.get("x1", ln.get("x", 0.0)))
                y = float(ln.get("y1", ln.get("y", 0.0)))
                text_for_line = gpt_lines[i] if i < len(gpt_lines) else ln.get("content", "")
                lines_for_layout.append({"text": text_for_line, "x": x, "y": y})
            pages_layout.append({
                "page_width": cached["pw"], "page_height": cached["ph"], "unit": "pixel",
                "lines": lines_for_layout
            })

            done += 1; progress.progress(done / total_to_process)

        del pages, page_numbers
        gc.collect()

    status_area.success("✅ すべてのページの処理が完了しました。")

    if all_corrected_texts:
        joined_txt = "\n\n".join(all_corrected_texts)
        st.download_button(
            "📥 補正テキストをダウンロード（TXT, ページ見出しなし）",
            data=joined_txt.encode("utf-8"), file_name="ocr_corrected.txt", mime="text/plain"
        )
    if pages_layout:
        try:
            st.download_button(
                "📥 Word（.docx：レイアウト近似）",
                data=build_docx_from_layout(pages_layout),
                file_name="ocr_layout.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            st.warning(f"Word出力に失敗しました：{e}")

# ===================== 画像：1ページ処理 =====================
else:
    # 画像として1ページ処理
    try:
        try:
            img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        except Exception:
            arr = np.frombuffer(file_bytes, dtype=np.uint8)
            bgr = cv2.imdecode(arr, cv2.IMREAD_COLOR)
            if bgr is None:
                raise ValueError("画像の読み込みに失敗しました（JPG/PNG/PDFのみ対応）。")
            img = Image.fromarray(cv2.cvtColor(bgr, cv2.COLOR_BGR2RGB))
        pages = [img]; page_numbers = [1]
    except Exception as e:
        st.exception(e); st.stop()

    st.caption("単一画像としてOCRします。")

    # --- 印影除去トグル（画像用） ---
    stamp_toggle_img = st.radio(
        "印影除去（ハンコ）", ["OFF", "ON"], index=0, horizontal=True, key="ui_stamp_toggle_img"
    )

    run_img = st.button("▶ この画像でOCRを実行", type="primary", key="run_img")
    if run_img:
        st.session_state["page_indices"] = [0]
        st.session_state["stamp_mode"] = stamp_toggle_img
    elif not st.session_state.get("page_indices"):
        # 初回の押し忘れでも1ページ実行にフォールバック
        st.session_state["page_indices"] = [0]
        st.session_state["stamp_mode"] = stamp_toggle_img

    stamp_mode = st.session_state.get("stamp_mode", stamp_toggle_img)

    all_corrected_texts: List[str] = []
    pages_layout: List[Dict[str, Any]] = []

    for page_img, page_num in zip(pages, page_numbers):
        st.write(f"## ページ {page_num}")

        orig_png = _bytes_from_pil(page_img)
        clean_img = remove_red_stamp(page_img) if stamp_mode == "ON" else page_img
        st.image(clean_img, caption=f"処理画像 (ページ {page_num})", use_container_width=True)

        buf = io.BytesIO(); clean_img.save(buf, format="PNG"); png_bytes = buf.getvalue()

        with st.spinner("OCRを実行中..."):
            t0 = time.perf_counter()
            try:
                cached = _ocr_dispatch(png_bytes, ocr_timeout, use_cache)
            except Exception as e:
                st.error(f"OCRに失敗：{e}")
                st.caption(f"OCR実行時間: {time.perf_counter() - t0:.1f}s")
                continue
            elapsed = time.perf_counter() - t0
            st.caption(f"OCR実行時間: {elapsed:.1f}s")

        azure_lines = cached.get("lines") or []
        default_text = "\n".join([ln["content"] for ln in azure_lines]) if azure_lines else (cached.get("raw") or "")
        if not default_text.strip():
            st.warning("OCRは成功しましたがテキストが空でした。DPIや画像品質を見直してください。")

        # 分類器（任意）で混同文字を自動置換
        dictionary = load_json_any(DICT_FILE)
        confusion_map = {k: v.get("right", v) if isinstance(v, dict) else v for k, v in dictionary.items()}
        auto_text = default_text
        repl_log = []
        if classifier_enable and cls_model is not None and labelmap is not None:
            auto_text, repl_log = auto_replace_with_classifier(
                page_img_raw=page_img,
                azure_lines=azure_lines,
                base_text=default_text,
                confusion_dict=confusion_map,
                cls_model=cls_model,
                labelmap=labelmap,
                conf_thresh=conf_thresh
            )
            if repl_log:
                st.info(f"🔁 自動置換 {len(repl_log)} 件（信頼度≥{conf_thresh:.2f}）")

        base_for_gpt = auto_text
        gpt_checked_text = base_for_gpt if skip_gpt else gpt_fix_text(base_for_gpt, dictionary)

        ocr_key = f"ocr_{page_num}"; gpt_key = f"gpt_{page_num}"; edit_key = f"edit_{page_num}"
        lines_key = f"lines_{page_num}"; imgraw_key = f"imgraw_{page_num}"

        if ocr_key not in st.session_state: st.session_state[ocr_key] = default_text
        if gpt_key not in st.session_state: st.session_state[gpt_key] = gpt_checked_text
        if edit_key not in st.session_state: st.session_state[edit_key] = gpt_checked_text
        st.session_state[lines_key] = azure_lines
        st.session_state[imgraw_key] = orig_png

        tab2, tab3, tab4 = st.tabs(["🖨️ OCRテキスト", "🤖 GPT/自動置換後", "✍️ 手作業修正"])
        with tab2:
            st.text_area(f"OCRテキスト（ページ {page_num}）", height=320, key=ocr_key)
            if repl_log:
                st.caption(f"自動置換ログ: {repl_log[:3]}{' ...' if len(repl_log)>3 else ''}")
        with tab3:
            st.text_area(f"GPT/自動置換後（ページ {page_num}）", height=320, key=gpt_key)
        with tab4:
            st.text_area(f"手作業修正（ページ {page_num}）", height=320, key=edit_key)
            if st.button(f"修正を保存 (ページ {page_num})", key=f"save_{page_num}"):
                corrected_text_current = st.session_state.get(edit_key, gpt_checked_text)
                learned = learn_charwise_with_missing(st.session_state.get(ocr_key, default_text), corrected_text_current)
                if learned:
                    update_dictionary_and_untrained(learned)
                    st.success(f"辞書と学習候補に {len(learned)} 件を追加しました！")
                else:
                    st.info("修正が検出されませんでした。")

                saved_cnt, err = save_evidence_assets(
                    backend_prefix=EVIDENCE_DIR,
                    page_num=page_num,
                    page_png_bytes_raw=st.session_state.get(imgraw_key, orig_png),
                    azure_lines=st.session_state.get(lines_key, azure_lines),
                    ocr_text=st.session_state.get(ocr_key, default_text),
                    corrected_text=corrected_text_current,
                    evidence_on=evidence_on
                )
                if err:
                    st.warning(f"学習証拠の保存に失敗：{err}")
                else:
                    st.success(f"学習証拠を保存しました（{saved_cnt}ファイル）")

        final_text_page = (st.session_state.get(edit_key) or st.session_state.get(gpt_key) or gpt_checked_text or default_text).strip()
        all_corrected_texts.append(final_text_page)

        gpt_lines = final_text_page.splitlines()
        lines_for_layout = []
        for i, ln in enumerate(azure_lines):
            x = float(ln.get("x1", ln.get("x", 0.0)))
            y = float(ln.get("y1", ln.get("y", 0.0)))
            text_for_line = gpt_lines[i] if i < len(gpt_lines) else ln.get("content", "")
            lines_for_layout.append({"text": text_for_line, "x": x, "y": y})
        pages_layout.append({"page_width": cached["pw"], "page_height": cached["ph"], "unit": "pixel", "lines": lines_for_layout})

    if all_corrected_texts:
        joined_txt = "\n\n".join(all_corrected_texts)
        st.download_button("📥 補正テキストをダウンロード（TXT, ページ見出しなし）",
                           data=joined_txt.encode("utf-8"), file_name="ocr_corrected.txt", mime="text/plain")
    if pages_layout:
        try:
            st.download_button("📥 Word（.docx：レイアウト近似）",
                               data=build_docx_from_layout(pages_layout),
                               file_name="ocr_layout.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        except Exception as e:
            st.warning(f"Word出力に失敗しました：{e}")
